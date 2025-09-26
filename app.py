"""
API Flask para transcripción de audio a PDF
Utiliza OpenAI Whisper para la transcripción y python-docx/reportlab para generar PDFs
"""

import os
import subprocess
import sys
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import whisper
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import inch
import datetime
import uuid
import threading
import queue as queue_module
import logging
from dotenv import load_dotenv

# Cargar variables de entorno desde .env
load_dotenv()

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# Leer modelo desde variable de entorno o usar 'base' por defecto
MODEL = os.environ.get("WHISPER_MODEL", "base")  # tiny, base, small, medium, large
ENABLE_ASYNC = os.environ.get("ENABLE_ASYNC", "1") == "1"

# Idiomas soportados (añade más si lo deseas)
SUPPORTED_LANGUAGES = [
    'spanish', 'english', 'portuguese', 'french', 'german', 'italian', 'auto'
]

def check_ffmpeg():
    """Verificar si FFmpeg está disponible"""
    try:
        subprocess.run(['ffmpeg', '-version'], capture_output=True, check=True)
        logger.info("FFmpeg encontrado en el sistema")
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        logger.warning("FFmpeg no encontrado")
        return False

def install_ffmpeg_windows():
    """Instalar FFmpeg en Windows usando pip"""
    try:
        logger.info("Intentando instalar ffmpeg-python...")
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'ffmpeg-python'])
        logger.info("ffmpeg-python instalado exitosamente")
        return True
    except subprocess.CalledProcessError as e:
        logger.error(f"Error instalando ffmpeg-python: {e}")
        return False

def setup_ffmpeg():
    """Configurar FFmpeg para el sistema"""
    if not check_ffmpeg():
        logger.info("FFmpeg no está disponible. Configurando alternativas...")
        
        # Intentar instalar ffmpeg-python como alternativa
        if install_ffmpeg_windows():
            logger.info("Configuración de FFmpeg completada")
        else:
            logger.warning("No se pudo configurar FFmpeg automáticamente")
            logger.info("Para mejor rendimiento, instala FFmpeg manualmente:")
            logger.info("1. Descarga desde https://ffmpeg.org/download.html")
            logger.info("2. O usa: winget install ffmpeg")
            logger.info("3. O usa: choco install ffmpeg")
    
    return True

app = Flask(__name__)

# Configurar CORS para permitir peticiones desde cualquier origen
CORS(app, resources={
    r"/health": {"origins": "*"},
    r"/api": {"origins": "*"},
    r"/transcribe": {"origins": "*"}
})

# Configuración
app.config['MAX_CONTENT_LENGTH'] = 1024 * 1024 * 1024  # 1GB máximo
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Crear directorios si no existen
for folder in ['uploads', 'outputs']:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Configurar FFmpeg
setup_ffmpeg()

#############################################
# CARGA DIFERIDA (LAZY LOAD) DEL MODELO    #
#############################################
# Evita cargar el modelo al iniciar el proceso, reduciendo uso de memoria inicial
# y tiempo de arranque en Render. Solo se cargará en la primera transcripción.

model = None  # se llenará bajo demanda

FALLBACK_ORDER = [
    MODEL,          # Modelo solicitado por el usuario
    "base",        # Fallback razonable
    "small",       # Más rápido que medium, mejor que tiny
    "tiny"         # Último recurso (rápido, menos precisión)
]

def load_whisper_model():
    """Cargar el modelo Whisper bajo demanda con degradación automática.
    Usa la variable de entorno WHISPER_AUTO_DOWNGRADE=1 para permitir bajar de modelo
    si hay errores de memoria o descarga.
    """
    global model
    if model is not None:
        return model

    auto_downgrade = os.environ.get("WHISPER_AUTO_DOWNGRADE", "1") == "1"
    tried = set()
    errors = {}

    logger.info("[Whisper] Carga diferida iniciada. Modelo solicitado: %s", MODEL)
    for candidate in FALLBACK_ORDER:
        if candidate in tried:
            continue
        tried.add(candidate)
        try:
            logger.info("[Whisper] Intentando cargar modelo: %s", candidate)
            m = whisper.load_model(candidate)
            model = m
            logger.info("[Whisper] Modelo '%s' cargado correctamente", candidate)
            if candidate != MODEL:
                logger.warning("[Whisper] Se degradó el modelo solicitado '%s' -> '%s'", MODEL, candidate)
            return model
        except Exception as e:
            errors[candidate] = str(e)
            logger.error("[Whisper] Error cargando '%s': %s", candidate, e)
            # Si no se permite auto downgrade, romper tras el primer fallo
            if not auto_downgrade:
                break

    # Si llegamos aquí, todos fallaron
    summary = "; ".join([f"{k}: {v[:120]}" for k,v in errors.items()])
    raise RuntimeError(f"No se pudo cargar ningún modelo Whisper. Errores: {summary}")

# Formatos de audio permitidos
ALLOWED_EXTENSIONS = {
    'mp3', 'wav', 'flac', 'm4a', 'aac', 'ogg', 'wma', 
    'mp4', 'avi', 'mov', 'mkv', 'webm'
}

def allowed_file(filename):
    """Verificar si el archivo tiene una extensión permitida"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def transcribe_audio(audio_path, language='spanish', progress_callback=None):
    """Transcribir archivo de audio usando Whisper.

    language: idioma forzado (si es 'auto' se deja que Whisper detecte).
    progress_callback: función opcional progress_callback(percentage:int, meta:dict)
    """
    try:
        # Verificar que el archivo existe
        if not os.path.exists(audio_path):
            raise FileNotFoundError(f"Archivo no encontrado: {audio_path}")
        
        # Verificar el tamaño del archivo
        file_size = os.path.getsize(audio_path)
        logger.info(f"Archivo a transcribir: {audio_path} ({file_size} bytes)")
        
        # Verificar permisos de lectura
        if not os.access(audio_path, os.R_OK):
            raise PermissionError(f"Sin permisos de lectura para: {audio_path}")
        
        logger.info(f"Iniciando transcripción de: {audio_path}")
        
        # Verificar si es un archivo WAV (formato nativo soportado)
        file_ext = os.path.splitext(audio_path)[1].lower()
        
        # Cargar modelo ahora (lazy)
        active_model = load_whisper_model()

        # Opciones comunes Whisper (forzamos español)
        base_options = {
            "fp16": False,
            "task": "transcribe",
            "verbose": False,
            "temperature": 0.0,
        }
        # Forzar idioma salvo que sea 'auto'
        if language and language.lower() != 'auto':
            base_options["language"] = language.lower()

        # Soporte opcional de CHUNKING para audios largos
        # Establecer CHUNK_SECONDS en entorno (por ejemplo 600 = 10 min) para habilitar.
        chunk_seconds_env = os.environ.get("CHUNK_SECONDS", "0")
        try:
            CHUNK_SECONDS = int(chunk_seconds_env)
        except ValueError:
            CHUNK_SECONDS = 0

        def _transcribe_wav_internal(wav_path):
            return active_model.transcribe(wav_path, **base_options)

        if file_ext == '.wav':
            # Archivo WAV: usar directamente
            logger.info("Archivo WAV detectado, procesando directamente")
            
            # Si chunking está activado y el archivo es largo, dividir
            if CHUNK_SECONDS > 0:
                try:
                    from pydub import AudioSegment
                    seg_audio = AudioSegment.from_wav(audio_path)
                    duration_sec = len(seg_audio) / 1000.0
                    if duration_sec > CHUNK_SECONDS:
                        logger.info(f"Aplicando chunking: duración {duration_sec:.1f}s > {CHUNK_SECONDS}s")
                        texts = []
                        segments_all = []
                        start_ms = 0
                        idx = 0
                        total_chunks = int((len(seg_audio) + (CHUNK_SECONDS*1000 - 1)) // (CHUNK_SECONDS*1000)) or 1
                        while start_ms < len(seg_audio):
                            end_ms = min(start_ms + CHUNK_SECONDS * 1000, len(seg_audio))
                            chunk = seg_audio[start_ms:end_ms]
                            temp_chunk = audio_path.replace('.wav', f'_chunk{idx}.wav')
                            chunk.export(temp_chunk, format='wav')
                            logger.info(f"Transcribiendo chunk {idx} ({start_ms/1000:.1f}-{end_ms/1000:.1f}s)...")
                            partial = _transcribe_wav_internal(temp_chunk)
                            texts.append(partial.get('text',''))
                            segs = partial.get('segments', [])
                            # Ajustar tiempos sumando offset
                            for s in segs:
                                if 'start' in s:
                                    s['start'] += start_ms/1000.0
                                if 'end' in s:
                                    s['end'] += start_ms/1000.0
                            segments_all.extend(segs)
                            try:
                                os.remove(temp_chunk)
                            except: pass
                            start_ms = end_ms
                            if progress_callback:
                                try:
                                    percentage = int(((idx + 1) / total_chunks) * 100)
                                    progress_callback(min(percentage, 99), {"chunk": idx+1, "total_chunks": total_chunks})
                                except Exception:
                                    pass
                            idx += 1
                        result = {"text": '\n'.join(texts), "segments": segments_all}
                    else:
                        result = _transcribe_wav_internal(audio_path)
                except ImportError:
                    logger.warning("Chunking requiere pydub instalado. Continuando sin dividir.")
                    result = _transcribe_wav_internal(audio_path)
                except Exception as ce:
                    logger.warning(f"Fallo chunking ({ce}); usando transcripción directa.")
                    result = _transcribe_wav_internal(audio_path)
            else:
                result = _transcribe_wav_internal(audio_path)
            
        else:
            # Para otros formatos, intentar conversión usando pydub si está disponible
            logger.info(f"Archivo {file_ext} detectado, intentando conversión")
            
            try:
                import pydub
                from pydub import AudioSegment
                
                # Cargar archivo con pydub
                audio = AudioSegment.from_file(audio_path)
                
                # Convertir a WAV temporal
                temp_wav = audio_path.replace(file_ext, '_temp.wav')
                audio.export(temp_wav, format="wav")
                
                logger.info(f"Archivo convertido a WAV temporal: {temp_wav}")
                
                # Transcribir el archivo WAV con español
                result = active_model.transcribe(temp_wav, **base_options)
                
                # Limpiar archivo temporal
                try:
                    os.remove(temp_wav)
                    logger.info("Archivo WAV temporal eliminado")
                except:
                    pass
                    
            except ImportError:
                logger.warning("pydub no está disponible, intentando conversión con FFmpeg")
                
                # Intentar conversión usando FFmpeg directamente
                try:
                    import subprocess
                    temp_wav = audio_path.replace(file_ext, '_temp.wav')
                    
                    # Intentar con ffmpeg si está disponible
                    try:
                        cmd = ['ffmpeg', '-i', audio_path, '-ar', '16000', '-ac', '1', '-y', temp_wav]
                        subprocess.run(cmd, capture_output=True, check=True)
                        logger.info(f"Archivo convertido con FFmpeg: {temp_wav}")
                        
                        # Transcribir el archivo WAV con español
                        result = active_model.transcribe(temp_wav, **base_options)
                        
                        # Limpiar archivo temporal
                        try:
                            os.remove(temp_wav)
                            logger.info("Archivo WAV temporal eliminado")
                        except:
                            pass
                            
                    except (subprocess.CalledProcessError, FileNotFoundError):
                        logger.warning("FFmpeg no disponible, intentando transcripción directa")
                        
                        # Intentar transcripción directa en español
                        result = active_model.transcribe(audio_path, **base_options)
                        
                except Exception as direct_error:
                    logger.error(f"Error en transcripción: {direct_error}")
                    raise Exception(
                        f"No se pudo procesar el archivo {file_ext}. "
                        f"Whisper con FFmpeg puede manejar la mayoría de formatos directamente. "
                        f"Error específico: {str(direct_error)}"
                    )
        
        if not result or "text" not in result:
            raise ValueError("La transcripción no produjo resultados válidos")
        
        # Verificar que hay contenido transcrito
        text = result["text"].strip()
        if not text:
            logger.warning("La transcripción está vacía - puede ser un archivo sin contenido de voz en español")
            # En lugar de fallar, devolver un mensaje informativo
            text = "[No se detectó contenido de voz en español en el archivo de audio]"
        
        return text, result.get("segments", [])
        
    except FileNotFoundError as e:
        logger.error(f"Archivo no encontrado: {str(e)}")
        raise e
    except PermissionError as e:
        logger.error(f"Error de permisos: {str(e)}")
        raise e
    except Exception as e:
        logger.error(f"Error en transcripción: {str(e)}")
        logger.error(f"Tipo de error: {type(e).__name__}")
        
        # Información adicional para debug
        if os.path.exists(audio_path):
            logger.error(f"Archivo existe: {audio_path}")
            logger.error(f"Tamaño: {os.path.getsize(audio_path)} bytes")
        else:
            logger.error(f"Archivo NO existe: {audio_path}")
            
        raise e

def create_pdf_with_reportlab(text, output_path, metadata=None):
    """Crear PDF usando ReportLab"""
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Crear estilos personalizados
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=30,
            alignment=TA_LEFT
        )
        
        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['Normal'],
            fontSize=11,
            textColor='black',
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leftIndent=0,
            rightIndent=0
        )
        
        story = []
        
        # Título
        title = Paragraph("Transcripción de Audio en Español", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Metadatos
        if metadata:
            meta_style = ParagraphStyle(
                'MetaData',
                parent=styles['Normal'],
                fontSize=10,
                textColor='grey',
                spaceAfter=20
            )
            
            meta_text = f"Fecha de transcripción: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}<br/>"
            if 'filename' in metadata:
                meta_text += f"Archivo original: {metadata['filename']}<br/>"
            if 'duration' in metadata:
                meta_text += f"Duración: {metadata['duration']:.2f} segundos<br/>"
            
            story.append(Paragraph(meta_text, meta_style))
            story.append(Spacer(1, 20))
        
        # Dividir el texto en párrafos para mejor legibilidad
        paragraphs = text.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                # Escapar caracteres especiales para ReportLab
                escaped_text = paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                p = Paragraph(escaped_text, content_style)
                story.append(p)
                story.append(Spacer(1, 6))
        
        doc.build(story)
        logger.info(f"PDF creado exitosamente: {output_path}")
        
    except Exception as e:
        logger.error(f"Error creando PDF: {str(e)}")
        raise e

def create_docx(text, output_path, metadata=None):
    """Crear documento DOCX usando python-docx"""
    try:
        doc = Document()
        
        # Título
        title = doc.add_heading('Transcripción de Audio en Español', 0)
        
        # Metadatos
        if metadata:
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.add_run(f"Fecha de transcripción: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n").italic = True
            if 'filename' in metadata:
                meta_paragraph.add_run(f"Archivo original: {metadata['filename']}\n").italic = True
            if 'duration' in metadata:
                meta_paragraph.add_run(f"Duración: {metadata['duration']:.2f} segundos\n").italic = True
            
            doc.add_paragraph()  # Espacio
        
        # Contenido
        doc.add_heading('Transcripción:', level=1)
        
        # Dividir texto en párrafos
        paragraphs = text.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc.save(output_path)
        logger.info(f"DOCX creado exitosamente: {output_path}")
        
    except Exception as e:
        logger.error(f"Error creando DOCX: {str(e)}")
        raise e

@app.route('/', methods=['GET'])
def index():
    """Endpoint principal - sirve la página web"""
    try:
        return send_file('index.html')
    except FileNotFoundError:
        return jsonify({'error': 'index.html no encontrado'}), 404
    except Exception as e:
        logger.error(f"Error sirviendo index: {e}")
        return jsonify({'error': 'Error interno sirviendo index'}), 500

@app.route('/api', methods=['GET'])
@app.route('/info', methods=['GET'])
def get_api_info():
    """Endpoint de información de la API"""
    loaded_name = None
    try:
        if model is not None and hasattr(model, 'dims') and hasattr(model.dims, 'n_vocab'):
            vocab = model.dims.n_vocab
            mapping = {
                51864: 'tiny',
                51865: 'base',
                51866: 'small',
                51867: 'medium'
            }
            loaded_name = mapping.get(vocab, 'large')
    except Exception:
        pass

    if loaded_name is None:
        model_info = f"(pendiente de carga lazy) solicitado={MODEL}"
    else:
        descriptors = {
            'tiny': 'rápido',
            'base': 'alta precisión',
            'small': 'muy alta precisión',
            'medium': 'excelente precisión',
            'large': 'máxima precisión'
        }
        model_info = f"{loaded_name} ({descriptors.get(loaded_name,'')})"

    return jsonify({
        'message': 'API de Transcripción de Audio a Documento - Optimizada para Español',
        'version': '1.0',
        'language_default': 'spanish',
        'model_requested': MODEL,
        'model_active': model_info,
        'lazy_loaded': model is not None,
        'languages_supported': SUPPORTED_LANGUAGES,
        'queue_enabled': ENABLE_ASYNC,
        'chunk_seconds': int(os.environ.get('CHUNK_SECONDS', '0') or 0),
        'endpoints': {
            'GET /': 'Página principal web',
            'GET /api': 'Información de la API (JSON)',
            'GET /health': 'Verificar estado de la API',
            'POST /transcribe': 'Transcribir archivo (sin cola, espera la respuesta)',
            'POST /transcribe_async': 'Crear tarea de transcripción en cola (si habilitado)',
            'GET /jobs/<id>': 'Estado de una tarea asíncrona',
            'GET /jobs/<id>/download': 'Descargar resultado de tarea completa'
        },
        'supported_formats': list(ALLOWED_EXTENSIONS),
        'max_file_size': '1GB'
    })

@app.route('/health', methods=['GET'])
def health():
    """Endpoint de salud"""
    return jsonify({
        'status': 'OK',
        'timestamp': datetime.datetime.now().isoformat(),
        'model_loaded': model is not None,
        'requested_model': MODEL
    })

#############################################
# COLA ASÍNCRONA SIMPLE EN MEMORIA          #
#############################################
job_queue = None
jobs = {}
jobs_lock = threading.Lock()

class JobStatus:
    PENDING = 'pending'
    PROCESSING = 'processing'
    DONE = 'done'
    ERROR = 'error'

def worker_loop():
    while True:
        job = job_queue.get()
        if job is None:
            break
        job_id = job['id']
        with jobs_lock:
            jobs[job_id]['status'] = JobStatus.PROCESSING
            jobs[job_id]['progress'] = 1
        try:
            def _progress(pct, meta=None):
                with jobs_lock:
                    if job_id in jobs:
                        jobs[job_id]['progress'] = pct
                        if meta:
                            jobs[job_id]['meta'] = meta
            text, segments = transcribe_audio(job['input_path'], job['language'], progress_callback=_progress)
            metadata = {
                'filename': job['original_filename'],
                'timestamp': job['timestamp']
            }
            if segments:
                last_segment = segments[-1]
                if 'end' in last_segment:
                    metadata['duration'] = last_segment['end']
            output_filename = f"{job['timestamp']}_transcription.{job['format']}"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            if job['format'] == 'pdf':
                create_pdf_with_reportlab(text, output_path, metadata)
            else:
                create_docx(text, output_path, metadata)
            with jobs_lock:
                jobs[job_id]['status'] = JobStatus.DONE
                jobs[job_id]['progress'] = 100
                jobs[job_id]['output_path'] = output_path
                jobs[job_id]['download_name'] = output_filename
        except Exception as e:
            logger.error(f"Job {job_id} error: {e}")
            with jobs_lock:
                jobs[job_id]['status'] = JobStatus.ERROR
                jobs[job_id]['error'] = str(e)
        finally:
            # Limpieza del archivo de entrada
            try:
                if os.path.exists(job['input_path']):
                    os.remove(job['input_path'])
            except Exception:
                pass
            job_queue.task_done()

if ENABLE_ASYNC:
    job_queue = queue_module.Queue()
    threading.Thread(target=worker_loop, daemon=True).start()
    logger.info("Cola asíncrona iniciada")

@app.route('/transcribe', methods=['POST'])
def transcribe():
    """Endpoint principal para transcribir audio"""
    try:
        # Verificar si se envió un archivo
        if 'audio' not in request.files:
            return jsonify({'error': 'No se encontró archivo de audio'}), 400
        
        file = request.files['audio']
        
        if file.filename == '':
            return jsonify({'error': 'No se seleccionó archivo'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                'error': 'Formato de archivo no soportado',
                'supported_formats': list(ALLOWED_EXTENSIONS)
            }), 400
        
        # Obtener parámetros opcionales
        output_format = request.form.get('format', 'pdf').lower()
        language = request.form.get('lang', 'spanish').lower()
        if language not in SUPPORTED_LANGUAGES:
            language = 'spanish'
        if output_format not in ['pdf', 'docx']:
            return jsonify({'error': 'Formato de salida debe ser "pdf" o "docx"'}), 400
        
        # Guardar archivo temporal
        filename = secure_filename(file.filename)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{filename}")
        
        # Asegurar que el directorio existe
        os.makedirs(os.path.dirname(input_path), exist_ok=True)
        
        file.save(input_path)
        logger.info(f"Archivo guardado: {input_path}")
        
        # Verificar que el archivo se guardó correctamente
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"No se pudo guardar el archivo en: {input_path}")
        
        file_size = os.path.getsize(input_path)
        logger.info(f"Archivo guardado correctamente. Tamaño: {file_size} bytes")
        
        # Transcribir audio
        logger.info("Iniciando transcripción...")
        transcription_text, segments = transcribe_audio(input_path, language=language)
        
        # Verificar que hay contenido (pero permitir el mensaje de "no speech")
        if not transcription_text or transcription_text.strip() == "":
            return jsonify({'error': 'No se pudo transcribir el audio - archivo vacío o sin contenido de voz en español'}), 400
        
        logger.info(f"Transcripción completada. Longitud: {len(transcription_text)} caracteres")
        logger.info(f"Texto transcrito: {transcription_text[:100]}...")  # Primeros 100 caracteres para debug
        
        # Preparar metadatos
        metadata = {
            'filename': filename,
            'timestamp': timestamp
        }
        
        # Calcular duración si está disponible en los segmentos
        if segments:
            last_segment = segments[-1]
            if 'end' in last_segment:
                metadata['duration'] = last_segment['end']
        
        # Crear archivo de salida
        output_filename = f"{timestamp}_transcription.{output_format}"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        if output_format == 'pdf':
            create_pdf_with_reportlab(transcription_text, output_path, metadata)
        else:  # docx
            create_docx(transcription_text, output_path, metadata)
        
        # Limpiar archivo temporal
        try:
            if os.path.exists(input_path):
                os.remove(input_path)
                logger.info(f"Archivo temporal eliminado: {input_path}")
        except Exception as cleanup_error:
            logger.warning(f"No se pudo eliminar archivo temporal: {cleanup_error}")
        
        # Retornar archivo
        return send_file(
            output_path,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/pdf' if output_format == 'pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except FileNotFoundError as e:
        logger.error(f"Archivo no encontrado: {str(e)}")
        return jsonify({'error': f'Archivo no encontrado: {str(e)}'}), 400
    except PermissionError as e:
        logger.error(f"Error de permisos: {str(e)}")
        return jsonify({'error': f'Error de permisos de archivo: {str(e)}'}), 403
    except Exception as e:
        logger.error(f"Error en transcripción: {str(e)}")
        
        # Limpiar archivo temporal en caso de error
        try:
            if 'input_path' in locals() and os.path.exists(input_path):
                os.remove(input_path)
        except:
            pass
            
        # Proporcionar información más específica del error
        if "ffmpeg" in str(e).lower():
            return jsonify({
                'error': 'Error de FFmpeg. Por favor instala FFmpeg para procesar archivos de audio.',
                'details': str(e),
                'solution': 'Instala FFmpeg desde https://ffmpeg.org/download.html o usa: winget install ffmpeg'
            }), 500
        else:
            return jsonify({'error': f'Error interno del servidor: {str(e)}'}), 500

@app.route('/transcribe_async', methods=['POST'])
def transcribe_async():
    if not ENABLE_ASYNC:
        return jsonify({'error': 'Modo asíncrono deshabilitado'}), 400
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No se encontró archivo de audio'}), 400
        file = request.files['audio']
        if file.filename == '':
            return jsonify({'error': 'No se seleccionó archivo'}), 400
        if not allowed_file(file.filename):
            return jsonify({'error': 'Formato de archivo no soportado', 'supported_formats': list(ALLOWED_EXTENSIONS)}), 400
        output_format = request.form.get('format', 'pdf').lower()
        language = request.form.get('lang', 'spanish').lower()
        if language not in SUPPORTED_LANGUAGES:
            language = 'spanish'
        if output_format not in ['pdf', 'docx']:
            return jsonify({'error': 'Formato de salida debe ser "pdf" o "docx"'}), 400
        filename = secure_filename(file.filename)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{timestamp}_{filename}")
        os.makedirs(os.path.dirname(input_path), exist_ok=True)
        file.save(input_path)
        job_id = str(uuid.uuid4())
        job_record = {
            'id': job_id,
            'status': JobStatus.PENDING,
            'progress': 0,
            'format': output_format,
            'language': language,
            'input_path': input_path,
            'original_filename': filename,
            'timestamp': timestamp,
            'created_at': datetime.datetime.utcnow().isoformat()
        }
        with jobs_lock:
            jobs[job_id] = job_record
        job_queue.put(job_record)
        return jsonify({'job_id': job_id, 'status': 'queued'})
    except Exception as e:
        logger.error(f"Error creando tarea: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/jobs/<job_id>', methods=['GET'])
def job_status(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Job no encontrado'}), 404
    filtered = {k: v for k, v in job.items() if k not in ('input_path',)}
    return jsonify(filtered)

@app.route('/jobs/<job_id>/download', methods=['GET'])
def job_download(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Job no encontrado'}), 404
    if job['status'] != JobStatus.DONE:
        return jsonify({'error': 'Job no está listo'}), 400
    output_path = job.get('output_path')
    if not output_path or not os.path.exists(output_path):
        return jsonify({'error': 'Archivo no disponible'}), 500
    return send_file(
        output_path,
        as_attachment=True,
        download_name=job.get('download_name', os.path.basename(output_path))
    )

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'Archivo demasiado grande. Máximo permitido: 1GB'}), 413

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Error interno del servidor'}), 500

if __name__ == '__main__':
    print("Iniciando API de Transcripción de Audio en Español...")
    print("Idioma configurado: Español")
    print("Formatos soportados:", ALLOWED_EXTENSIONS)
    print("Tamaño máximo de archivo: 1GB")
    print("Modelo Whisper:", MODEL)
    print("\nEndpoints disponibles:")
    print("  GET  /           - Información de la API")
    print("  GET  /health     - Estado de la API")
    print("  POST /transcribe - Transcribir archivo de audio en español")
    print("\nEjemplo de uso:")
    print("  curl -X POST -F 'audio=@archivo.mp3' -F 'format=pdf' http://localhost:5000/transcribe")
    
    # Configuración específica para desarrollo local y Render
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('FLASK_ENV', 'production') == 'development'
    
    print(f"Iniciando servidor en puerto: {port}")
    print(f"Modo debug: {debug_mode}")
    
    # Solo ejecutar app.run() en desarrollo local
    # En producción (Render) se usa Gunicorn
    if os.environ.get('RENDER') != '1':
        app.run(debug=debug_mode, host='0.0.0.0', port=port)
else:
    # Para Gunicorn - configuración de producción
    print("Ejecutándose con Gunicorn en producción")
    print("Modelo Whisper cargado:", MODEL)
    print("CORS habilitado para todos los orígenes")
